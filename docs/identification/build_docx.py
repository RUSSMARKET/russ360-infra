#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Идентифицирующие материалы ПО «Russmarket 360».
Сборка единого документа: титул, содержание, листинги исходного кода,
экранные формы (код компонента + снимок экрана).

Запуск:  python3 build_docx.py
Выход:   ./russ360-identification.docx
"""

import os
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(HERE, "..", ".."))
SCREENS = os.path.join(HERE, "screens")
OUT = os.path.join(HERE, "russ360-identification.docx")
TOC_JSON = os.path.join(HERE, "toc.json")       # title -> page (заполняется 2-м проходом)
HEAD_JSON = os.path.join(HERE, "headings.json")  # упорядоченный список [level, title]

# ── Гарнитуры (метрически совместимые с MS Office) ──────────────────────────
F_BODY = "Times New Roman"
F_HEAD = "Arial"
F_CODE = "Courier New"

# ── Палитра: сдержанная, деловая ────────────────────────────────────────────
INK = RGBColor(0x1A, 0x1A, 0x1A)        # основной текст
NAVY = RGBColor(0x1F, 0x3A, 0x5F)       # заголовки
GREY = RGBColor(0x60, 0x6A, 0x78)       # подписи, комментарии
CODE_BG = "F4F5F7"
CODE_BORDER = "C9CFD8"
RULE = "B9C0CC"

IMG_WIDTH = Cm(15.2)

# счётчики подписей
_fig = [0]
_lst = [0]


def excerpt(relpath, nlines=100000, start=1):
    full = os.path.join(REPO, relpath)
    try:
        with open(full, encoding="utf-8") as f:
            lines = f.read().split("\n")
    except OSError:
        return "// исходный файл не найден: " + relpath
    chunk = lines[start - 1:start - 1 + nlines]
    text = "\n".join(chunk).rstrip("\n")
    if len(lines) > start - 1 + nlines:
        text += "\n    /* ... фрагмент исходного кода ... */"
    return text


# ── Данные ──────────────────────────────────────────────────────────────────

LISTINGS = [
    {"group": "Сервис аутентификации (rusaiauth)", "files": [
        {"t": "Формирование JWT с издателем (iss)",
         "p": "rusaiauth/app/Domain/Identity/OAuth/IssuerAwareAccessToken.php"},
    ]},
    {"group": "Доменное ядро (rusaicore)", "files": [
        {"t": "Доменная модель «Проект»", "p": "rusaicore/app/Domain/Project/Project.php"},
        {"t": "Перечисление статусов проекта", "p": "rusaicore/app/Domain/Project/ProjectStatus.php"},
        {"t": "Доменная модель «Участие в проекте»", "p": "rusaicore/app/Domain/ProjectMembership/ProjectMembership.php"},
        {"t": "Доменная модель «Сотрудник»", "p": "rusaicore/app/Domain/Employee/Employee.php"},
        {"t": "Доменная модель «Операционная точка»", "p": "rusaicore/app/Domain/OperationalLocation/OperationalLocation.php"},
        {"t": "Проверка маркера доступа (OAuth)", "p": "rusaicore/app/Infrastructure/Auth/OAuthTokenValidator.php"},
        {"t": "Создание участия в проекте", "p": "rusaicore/app/Application/ProjectMembership/Actions/CreateProjectMembership.php"},
        {"t": "Выборка участий в проекте", "p": "rusaicore/app/Application/ProjectMembership/Actions/ListProjectMemberships.php"},
        {"t": "Промежуточный обработчик: проверка прав", "p": "rusaicore/app/Http/Middleware/EnsureScope.php"},
        {"t": "Промежуточный обработчик: идемпотентность", "p": "rusaicore/app/Http/Middleware/EnsureIdempotency.php"},
    ]},
    {"group": "Сервис складского учёта (rusaisklad_back)", "files": [
        {"t": "Клиент обращения к ядру (HTTP)", "p": "rusaisklad_back/app/Domain/Core/Gateways/CoreApiClient.php"},
        {"t": "Шлюз проектов", "p": "rusaisklad_back/app/Domain/Core/Gateways/CoreProjectGateway.php"},
        {"t": "Шлюз сотрудников", "p": "rusaisklad_back/app/Domain/Core/Gateways/CoreEmployeeGateway.php"},
        {"t": "Шлюз участий в проекте", "p": "rusaisklad_back/app/Domain/Core/Gateways/CoreProjectMembershipGateway.php"},
        {"t": "Контракт каталога проектов", "p": "rusaisklad_back/app/Domain/Core/Contracts/ProjectCatalog.php"},
        {"t": "Контракт справочника сотрудников", "p": "rusaisklad_back/app/Domain/Core/Contracts/EmployeeDirectory.php"},
        {"t": "Объект передачи данных проекта", "p": "rusaisklad_back/app/Domain/Core/DTOs/CoreProjectData.php"},
        {"t": "Объект передачи данных участия", "p": "rusaisklad_back/app/Domain/Core/DTOs/CoreProjectMembershipData.php"},
    ]},
    {"group": "Веб-клиент «Полевые продажи» (fintech)", "files": [
        {"t": "Подготовка параметров входа (PKCE)", "p": "fintech/src/shared/lib/oidcPkce.ts"},
        {"t": "Обновление и хранение маркеров доступа", "p": "fintech/src/shared/lib/oidcRefresh.ts"},
        {"t": "Хранилище состояния сессии", "p": "fintech/src/shared/auth/useSessionStore.ts"},
        {"t": "Контроль доступа к маршрутам", "p": "fintech/middleware/auth.global.ts"},
        {"t": "Хранилище маркера доступа", "p": "fintech/src/shared/auth/tokenStorage.ts"},
        {"t": "Доступ к маркеру доступа", "p": "fintech/src/shared/auth/accessToken.ts"},
        {"t": "Восстановление пользователя и сессии", "p": "fintech/src/entities/user/model.ts", "lines": 132},
        {"t": "Клиент HTTP с обновлением маркера", "p": "fintech/src/shared/api/http.ts"},
        {"t": "Завершение сеанса (logout)", "p": "fintech/src/shared/lib/oidcLogout.ts"},
    ]},
    {"group": "Веб-клиент «Складской учёт» (rusaisklad_front)", "files": [
        {"t": "Шифрование локального хранилища", "p": "rusaisklad_front/src/shared/utils/encryption.ts"},
    ]},
]

SURFACES = [
    {"name": "Подсистема входа (единая аутентификация)",
     "base": "https://sso.rusaifin.ru/auth",
     "intro": "Единая страница входа для всех приложений системы. Авторизация выполняется "
              "по номеру телефона и паролю либо одноразовому коду. После успешного входа "
              "пользователь возвращается в выбранное приложение.",
     "pages": [
        {"title": "Страница входа", "route": "/auth/login",
         "path": "rusaiauth/resources/js/preview/pages/LoginPhonePage.vue",
         "screenshot": "auth-login-phone.png"},
     ]},
    {"name": "Приложение «Полевые продажи»",
     "base": "https://fintech.rusaifin.ru",
     "intro": "Рабочее место сотрудника полевых продаж: рабочий стол, каталог продуктов, "
              "сотрудники и проекты, планы и показатели, отчётность, заявки, задачи, "
              "обучение, оформление сотрудников и настройки учётной записи.",
     "pages": [
        {"title": "Рабочий стол", "route": "/", "path": "fintech/src/pages/home/ui/Home.vue", "screenshot": "fin-home.png"},
        {"title": "Продукты", "route": "/products", "path": "fintech/src/pages/products/ui/Products.vue", "screenshot": "fin-products.png"},
        {"title": "Сотрудники", "route": "/agents", "path": "fintech/src/pages/agents/ui/Agents.vue", "screenshot": "fin-agents.png"},
        {"title": "Проекты", "route": "/project", "path": "fintech/src/pages/project/ui/Project.vue", "screenshot": "fin-project.png"},
        {"title": "Планы и показатели", "route": "/kpi", "path": "fintech/src/pages/kpi/ui/Kpi.vue", "screenshot": "fin-kpi.png"},
        {"title": "Отчётность", "route": "/reporting", "path": "fintech/src/pages/reporting/ui/Reporting.vue", "screenshot": "fin-reporting.png"},
        {"title": "Заявки", "route": "/requests", "path": "fintech/src/pages/requests/ui/Requests.vue", "screenshot": "fin-requests.png"},
        {"title": "Задачи", "route": "/tasks", "path": "fintech/src/pages/tasks/ui/Tasks.vue", "screenshot": "fin-tasks.png"},
        {"title": "Инвентарь", "route": "/inventory", "path": "fintech/src/pages/inventory/ui/Inventory.vue", "screenshot": "fin-inventory.png"},
        {"title": "Обучение", "route": "/education", "path": "fintech/src/pages/education/ui/Education.vue", "screenshot": "fin-education.png"},
        {"title": "Оформление сотрудников", "route": "/onboarding", "path": "fintech/src/pages/onboarding/ui/Onboarding.vue", "screenshot": "fin-onboarding.png"},
        {"title": "Настройки", "route": "/options", "path": "fintech/src/pages/options/ui/Options.vue", "screenshot": "fin-options.png"},
        {"title": "Журнал действий", "route": "/logging", "path": "fintech/src/pages/logging/ui/Logging.vue", "screenshot": "fin-logging.png"},
        {"title": "Профиль", "route": "/profile", "path": "fintech/src/pages/profile/ui/Profile.vue", "screenshot": "fin-profile.png"},
     ]},
    {"name": "Приложение «Складской учёт»",
     "base": "https://rusaisklad.ru",
     "intro": "Рабочее место складского учёта: остатки товарно-материальных ценностей, "
              "номенклатура (SKU), передачи и движения ценностей, проекты, сотрудники и "
              "настройки учётной записи.",
     "pages": [
        {"title": "Рабочий стол", "route": "/", "path": "rusaisklad_front/src/pages/home/ui/Home.vue", "screenshot": "sklad-home.png"},
        {"title": "Остатки ТМЦ", "route": "/inventory", "path": "rusaisklad_front/src/pages/inventory_tmc/views/InventoryBalancesView.vue", "screenshot": "sklad-inventory.png"},
        {"title": "Номенклатура (SKU)", "route": "/sku", "path": "rusaisklad_front/src/pages/sku/ui/Sku.vue", "screenshot": "sklad-sku.png"},
        {"title": "Передачи ТМЦ", "route": "/reporting", "path": "rusaisklad_front/src/app/routes/reporting.vue", "screenshot": "sklad-reporting.png"},
        {"title": "Проекты", "route": "/project", "path": "rusaisklad_front/src/pages/project/ui/Project.vue", "screenshot": "sklad-project.png"},
        {"title": "Сотрудники", "route": "/users", "path": "rusaisklad_front/src/pages/users/ui/Agents.vue", "screenshot": "sklad-users.png"},
        {"title": "Профиль", "route": "/profile", "path": "rusaisklad_front/src/pages/profile/ui/Profile.vue", "screenshot": "sklad-profile.png"},
     ]},
]


# ── Низкоуровневые помощники XML ────────────────────────────────────────────

def _sub(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    el.append(e)
    return e


def keep_next(p):
    _sub(p._p.get_or_add_pPr(), "w:keepNext")


def keep_lines(p):
    _sub(p._p.get_or_add_pPr(), "w:keepLines")


def shade(p, fill):
    _sub(p._p.get_or_add_pPr(), "w:shd", **{"w:val": "clear", "w:fill": fill})


def set_run(run, font=F_BODY, size=11.5, color=INK, bold=False, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _sub(rpr, "w:rFonts")
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)
    return run


# ── Абзацы ──────────────────────────────────────────────────────────────────

def para(doc, space_after=6, space_before=0, align=None, line=1.15):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def body(doc, text, size=11.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=INK, space_after=8, first_line=True):
    p = para(doc, space_after=space_after, align=align)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.0)
    set_run(p.add_run(text), size=size, color=color)
    return p


def heading(doc, text, level):
    sizes = {1: 15, 2: 12.5, 3: 11.5}
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else (10 if level == 2 else 8))
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    p.paragraph_format.keep_with_next = True
    for r in list(p.runs):
        r.text = ""
    set_run(p.add_run(text), font=F_HEAD, size=sizes[level], color=NAVY, bold=True)
    return p


def caption(doc, text):
    p = para(doc, space_after=12, space_before=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run(text), size=9.5, color=GREY, italic=True)
    return p


def code_caption(doc, text):
    p = para(doc, space_after=2, space_before=8)
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size=9.5, color=GREY, italic=True)
    return p


def file_ref(doc, path):
    p = para(doc, space_after=3, space_before=0)
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run("Файл: "), size=9, color=GREY)
    set_run(p.add_run(path), font=F_CODE, size=8.5, color=GREY)
    return p


# ── Листинг кода: одна ячейка-таблица со светлой заливкой и тонкой рамкой ─────

def _split_comment(line):
    """Грубое деление строки на (код, комментарий) для приглушения комментариев."""
    s = line.lstrip()
    if s.startswith(("//", "/*", "*", "*/", "#")):
        return "", line
    idx = line.find("//")
    if idx >= 0 and line.count('"', 0, idx) % 2 == 0 and line.count("'", 0, idx) % 2 == 0:
        return line[:idx], line[idx:]
    return line, ""


def code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    _sub(tblPr, "w:tblW", **{"w:type": "dxa", "w:w": "9300"})
    _sub(tblPr, "w:tblLayout", **{"w:type": "fixed"})
    borders = _sub(tblPr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _sub(borders, "w:" + edge, **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": CODE_BORDER})
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    _sub(tcPr, "w:shd", **{"w:val": "clear", "w:fill": CODE_BG})
    mar = _sub(tcPr, "w:tcMar")
    for edge, w in (("top", "60"), ("bottom", "60"), ("left", "120"), ("right", "120")):
        _sub(mar, "w:" + edge, **{"w:type": "dxa", "w:w": w})

    lines = code.split("\n")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in lines:
        p = cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0
        keep_lines(p)
        code_part, comment_part = _split_comment(line if line else " ")
        if code_part:
            set_run(p.add_run(code_part), font=F_CODE, size=8.5, color=INK)
        if comment_part:
            set_run(p.add_run(comment_part), font=F_CODE, size=8.5, color=GREY, italic=True)
        if not code_part and not comment_part:
            set_run(p.add_run(" "), font=F_CODE, size=8.5, color=INK)
    para(doc, space_after=6, space_before=0)
    return table


def picture(doc, fname):
    path = os.path.join(SCREENS, fname)
    p = para(doc, space_after=2, space_before=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    if os.path.exists(path):
        run = p.add_run()
        run.add_picture(path, width=IMG_WIDTH)
    else:
        set_run(p.add_run("[ изображение отсутствует: %s ]" % fname), size=9, color=GREY, italic=True)
    return p


# ── Служебные поля (нумерация страниц, оглавление) ──────────────────────────

def _field(paragraph, instr, placeholder=""):
    r = paragraph.add_run()
    _sub(r._r, "w:fldChar", **{"w:fldCharType": "begin"})
    r2 = paragraph.add_run()
    it = _sub(r2._r, "w:instrText", **{"xml:space": "preserve"})
    it.text = instr
    r3 = paragraph.add_run()
    _sub(r3._r, "w:fldChar", **{"w:fldCharType": "separate"})
    r4 = paragraph.add_run(placeholder)
    r5 = paragraph.add_run()
    _sub(r5._r, "w:fldChar", **{"w:fldCharType": "end"})
    return [r, r2, r3, r4, r5]


def render_toc(doc, titles, pages, tab_pos_cm):
    """Реальное оглавление с отточием. pages: dict title->page (или None)."""
    for level, title in titles:
        p = para(doc, space_after=5 if level == 1 else 3, line=1.1)
        if level == 2:
            p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_pos_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        set_run(p.add_run(title), size=11.5 if level == 1 else 11, color=INK, bold=(level == 1))
        if pages and title in pages:
            set_run(p.add_run("\t" + str(pages[title])), size=11, color=INK, bold=(level == 1))


def setup_footer(section):
    section.different_first_page_header_footer = True
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runs = _field(fp, "PAGE", "")
    for r in runs:
        set_run(r, size=10, color=GREY)
    # верхний колонтитул — сдержанная строка с названием
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(hp.add_run("Russmarket 360"), size=8.5, color=GREY)
    _sub(hp._p.get_or_add_pPr(), "w:pBdr")
    pbdr = hp._p.get_or_add_pPr().find(qn("w:pBdr"))
    _sub(pbdr, "w:bottom", **{"w:val": "single", "w:sz": "4", "w:space": "4", "w:color": RULE})


def enable_update_fields(doc):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        _sub(settings, "w:updateFields", **{"w:val": "true"})


# ── Сборка ──────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Базовый стиль
    normal = doc.styles["Normal"]
    normal.font.name = F_BODY
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.15

    # Формат А4, поля близко к деловому стандарту
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    setup_footer(sec)

    # Состав оглавления (заголовки уровней 1–2) — единые строки для TOC и тела
    toc_titles = [(1, "Общие сведения"), (1, "1. Листинги исходного кода")]
    for i, grp in enumerate(LISTINGS, 1):
        toc_titles.append((2, "1.%d. %s" % (i, grp["group"])))
    toc_titles.append((1, "2. Экранные формы пользовательского интерфейса"))
    for i, surf in enumerate(SURFACES, 1):
        toc_titles.append((2, "2.%d. %s" % (i, surf["name"])))
    toc_pages = {}
    if os.path.exists(TOC_JSON):
        with open(TOC_JSON, encoding="utf-8") as f:
            toc_pages = json.load(f)
    toc_tab = (21.0 - 2.5 - 1.5)  # ширина текста, см

    # ── Титульный лист ──
    for _ in range(5):
        doc.add_paragraph()
    t = para(doc, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(t.add_run("ИДЕНТИФИЦИРУЮЩИЕ МАТЕРИАЛЫ"), font=F_HEAD, size=20, color=NAVY, bold=True)
    t2 = para(doc, space_after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(t2.add_run("программного обеспечения"), font=F_HEAD, size=14, color=INK)
    t3 = para(doc, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(t3.add_run("«Russmarket 360»"), font=F_HEAD, size=22, color=NAVY, bold=True)

    rule = para(doc, space_after=16, space_before=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    pbdr = _sub(rule._p.get_or_add_pPr(), "w:pBdr")
    _sub(pbdr, "w:bottom", **{"w:val": "single", "w:sz": "6", "w:space": "1", "w:color": RULE})

    d = para(doc, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(d.add_run("Программный код и экранные формы пользовательского интерфейса"),
            size=12, color=INK)
    for _ in range(10):
        doc.add_paragraph()
    yr = para(doc, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(yr.add_run("2026"), size=12, color=INK)
    doc.add_page_break()

    # ── Содержание ──
    heading(doc, "Содержание", 1)
    render_toc(doc, toc_titles, toc_pages, toc_tab)
    doc.add_page_break()

    # ── Введение ──
    heading(doc, "Общие сведения", 1)
    body(doc,
         "Программное обеспечение «Russmarket 360» представляет собой совокупность "
         "серверных приложений и веб-клиентов, объединённых единой подсистемой "
         "аутентификации. Серверная часть реализована на платформе Laravel (язык PHP), "
         "клиентская — на платформе Nuxt (язык TypeScript, библиотека Vue). Взаимодействие "
         "приложений и разграничение доступа построено на протоколах OAuth 2.0 и OpenID "
         "Connect.")
    body(doc,
         "Настоящий документ содержит идентифицирующие материалы: листинги исходного кода "
         "ключевых модулей программы (раздел 1) и экранные формы пользовательского "
         "интерфейса с указанием соответствующих компонентов исходного кода (раздел 2). "
         "Снимки экранов получены на испытательном стенде; используемые данные обезличены.")

    # ── Раздел 1. Листинги ──
    heading(doc, "1. Листинги исходного кода", 1)
    body(doc,
         "Ниже приведены листинги исходного кода ключевых модулей по составным частям "
         "программного обеспечения: доменная модель, подсистема аутентификации, средства "
         "межсервисного взаимодействия и клиентский слой управления сессией.")
    sub = 0
    for grp in LISTINGS:
        sub += 1
        heading(doc, "1.%d. %s" % (sub, grp["group"]), 2)
        for item in grp["files"]:
            _lst[0] += 1
            code_caption(doc, "Листинг %d — %s" % (_lst[0], item["t"]))
            file_ref(doc, item["p"])
            code_block(doc, excerpt(item["p"], item.get("lines", 100000)))

    # ── Раздел 2. Экранные формы ──
    doc.add_page_break()
    heading(doc, "2. Экранные формы пользовательского интерфейса", 1)
    body(doc,
         "Для каждой экранной формы приведён фрагмент исходного кода соответствующего "
         "компонента интерфейса и снимок экрана с её внешним видом.")
    sub = 0
    for surf in SURFACES:
        sub += 1
        heading(doc, "2.%d. %s" % (sub, surf["name"]), 2)
        ap = para(doc, space_after=4)
        set_run(ap.add_run("Веб-адрес: "), size=10, color=GREY)
        set_run(ap.add_run(surf["base"]), font=F_CODE, size=9, color=GREY)
        if surf.get("intro"):
            body(doc, surf["intro"], size=11, space_after=8)
        for pg in surf["pages"]:
            _lst[0] += 1
            _fig[0] += 1
            h = para(doc, space_after=2, space_before=8)
            h.paragraph_format.keep_with_next = True
            set_run(h.add_run("%s" % pg["title"]), font=F_HEAD, size=11, color=NAVY, bold=True)
            set_run(h.add_run("   (%s)" % pg["route"]), font=F_CODE, size=9, color=GREY)
            code_caption(doc, "Листинг %d — компонент экранной формы" % _lst[0])
            file_ref(doc, pg["path"])
            code_block(doc, excerpt(pg["path"], 28))
            picture(doc, pg["screenshot"])
            caption(doc, "Рисунок %d — экранная форма «%s»" % (_fig[0], pg["title"]))

    doc.save(OUT)
    with open(HEAD_JSON, "w", encoding="utf-8") as f:
        json.dump(toc_titles, f, ensure_ascii=False)
    missing = [pg["screenshot"] for s in SURFACES for pg in s["pages"]
               if not os.path.exists(os.path.join(SCREENS, pg["screenshot"]))]
    print("OK ->", OUT)
    print("Листингов:", _lst[0], "| рисунков:", _fig[0],
          "| оглавление:", "со страницами" if toc_pages else "без страниц (1-й проход)")
    print("Скриншоты: все на месте." if not missing else ("Отсутствуют: " + ", ".join(missing)))


if __name__ == "__main__":
    main()
